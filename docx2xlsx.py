#!/usr/bin/env python3

import json
import os
import sys
import optparse

import xlsxwriter
from docx import Document

prog_VERSION = "0.1.0-dev"

def parse_options(args=None, values=None):
    """
    Define and parse `optparse` options for command-line usage.
    """
    usage = """%prog -d docxs_in_FOLDER -o OUTPUT_FILE-xlsx -c CONFIG_FILE """
    desc = "reading docx files in the specified folder and write the information to Excel, according to the config files"
    ver = "%%prog %s" % prog_VERSION

    parser = optparse.OptionParser(usage=usage, description=desc, version=ver)

    parser.add_option("-d", "--dir", dest="docxs_dir", default=None,
                      help="specify the Markdown file *.md which will be transfer to html")
    parser.add_option("-o", "--output", dest="output_file", default=None,
                      help="specify the output *.html file which will be transfer to html")
    parser.add_option("-c", "--config", dest="config_file", default=None,
                      help="specify the config file which defines the rules how to subsitute the tags")

    (options, args) = parser.parse_args(args, values)
    # print(options)

    for arg in vars(options):
        filename = getattr(options, arg) # not completing the options 
        if filename is None:
            parser.print_usage()
            sys.exit(1)
        if (arg == "config_file" and not os.path.isfile(filename)) or (arg == "docxs_dir" and not os.path.isdir(filename)):
            print("invalid file name specified : %s %s " % (arg, filename))
            sys.exit(2)

    return options

def parse_range(numbers, end=-1) :
    # 1~3 : [1,2,3]
    # 1,+2 : base number is 1, and step with 2 => 1,3,5,..., <=end
    # 3,5,10: [3,5,10]
    if '~' in numbers:
        i = numbers.index('~')
        fro = int(numbers[:i])
        to = int(numbers[i+1:])
        return list(range(fro, to+1))
    elif ',' in numbers:
        if '+' in numbers:
            assert end >= 1, "you should set 'end' paramenter value when parsing string such as '1,+2'"
            i = numbers.index(',')
            base = int(numbers[:i])
            i = numbers.index('+')
            step = int(numbers[i+1:])

            data = []
            while base <= end :
                data.append(base)
                base += step

            return data
        else :
            # it should be something like "3,5,8,10"
            data = []
            for num in numbers.split(','):
                data.append(int(num))
            return data
    else :
        return [int(numbers)]
        
def parse_config_file(cfg_file) :
    with open(cfg_file, "r") as f:
        cfg_data = json.load(f)
        f.close()
        return cfg_data

def get_the_exactly_table(tables):
    # if len(tables) == 1:
    #     return tables[0]

    for t in tables :
        txt = t.cell(1,0).text
        txt = "".join(txt.split())
        if txt == '姓名' :
            return t

    return None

def get_clean_text(cell, row, col) :
    txt = cell(row, col).text
    return "".join(txt.split())

def transfer(docxs_dir, output_file, config_file):
    cfg_data = parse_config_file(config_file)

    #Create a workbook and add a worksheet.
    workbook = xlsxwriter.Workbook(output_file)
    worksheet = workbook.add_worksheet()
    out_row = 1 # row 0 for title
    out_col = 0
    out_row0_done = False
    cell_format = workbook.add_format()
    cell_format.set_text_wrap()
    cell_format.set_align('center')
    cell_format.set_align('vcenter')

    # cell_format.set_align('vjustify')
    # print(cfg_data)
    xlsx_row_title = 1
    xlsx_row = xlsx_row_title + 1
    for file in os.listdir(docxs_dir):
        filename = os.fsdecode(file)

        title_has_wrote = False
        if filename.endswith(".docx") or filename.endswith(".doc"):
            path_file = os.path.join(docxs_dir, filename)
            document = Document(path_file)
            table = get_the_exactly_table(document.tables)
            if not table :
                print("invalid table in docx file : %s" % path_file)
                exit(2)

            for section in cfg_data :
                r = cfg_data.get(section)
                docx_row = r.get('row')
                # print("get the row: %d" % r.get('row'))
                fro = r.get('from')
                docx_cols = parse_range(fro.get('col'))
                end = docx_cols[-1]
                docx_keys = parse_range(fro.get('key'), end)
                docx_values = parse_range(fro.get('val'), end)
                to = r.get('to')

                xlsx_col = to.get('col_start')
                if not title_has_wrote : 
                    for c in docx_keys :
                        txt = get_clean_text(table.cell, docx_row, c)
                        # print("%s:" % txt)
                        worksheet.write(xlsx_row_title, xlsx_col, txt, cell_format)
                        xlsx_col += 1

                xlsx_col = to.get('col_start')
                for c in docx_values :
                    txt = get_clean_text(table.cell, docx_row, c)
                    # print("[%s]" % txt)
                    worksheet.write(xlsx_row, xlsx_col, txt, cell_format)
                    xlsx_col += 1

            xlsx_row += 1
            if not title_has_wrote :
                title_has_wrote = True
            
    workbook.close()


if __name__ == "__main__" :
    
    opts = parse_options()
    transfer(opts.docxs_dir, opts.output_file, opts.config_file)

# Load the first table from your document. In your example file,
# there is only one table, so I just grab the first one.
# document = Document('test.docx')
# table = document.tables[0]


# Create a workbook and add a worksheet.
# workbook = xlsxwriter.Workbook(os.path.join(directory, '九三学社青工委委员推荐表.xlsx'))
# worksheet = workbook.add_worksheet()
# out_row = 1 # row 0 for title
# out_col = 0
# out_row0_done = False
# cell_format = workbook.add_format()
# cell_format.set_text_wrap()
# cell_format.set_align('center')
# cell_format.set_align('vcenter')
# # cell_format.set_align('vjustify')

# for file in os.listdir(directory):
#     filename = os.fsdecode(file)
#     if filename.endswith(".docx"):
#         print(os.path.join(directory, filename))
#         document = Document(os.path.join(directory, filename))
#         table = document.tables[0]

#         count = len(table.rows)

#         if count == 11 :

#             if not out_row0_done:
#                 out_col = 0
#                 for v in keys :
#                     txt = table.cell(v[0], v[1]).text
#                     txt = "".join(txt.split())
#                     cell_width = len(txt)

#                     worksheet.set_column(0, out_col, cell_width)
#                     worksheet.write(0, out_col, txt, cell_format)
#                     out_col += 1
#                 out_row0_done = True

#                 # print("[%s] \n" % txt)

#             out_col = 0
#             for v in vals :
#                 txt = table.cell(v[0], v[1]).text.strip()
#                 worksheet.write(out_row, out_col, txt, cell_format)

#                 # print("[%s] \n" % txt)
#                 out_col += 1

#             out_row += 1

#         elif count == 10 :
#             out_col = 0
#             wrote_col_10 = False
#             for v in val2s :
#                 txt = table.cell(v[0], v[1]).text.strip()

#                 # if v[0] == politicalparty_row and v[1] == politicalparty_col :
#                 #     out_col += 1
#                 # elif v[0] == join_row and v[1]==join_col :
#                 #     worksheet.write(out_row, out_col + 3, txt)
#                 #     continue
#                 # elif v[0] == specialtitle_row and v[1] == specialtitle_col :
#                 #     out_col += 2

#                 if out_col == 6:
#                     out_col += 1
#                 elif out_col == 8 :
#                     if not wrote_col_10:
#                         worksheet.write(out_row, 10, txt, cell_format)
#                         wrote_col_10 = True
#                         continue
#                 elif out_col == 9:
#                     out_col += 3

#                 worksheet.write(out_row, out_col, txt, cell_format)

#                 # print("[%s] \n" % txt)
#                 out_col += 1

#             out_row += 1

#         else :

#             print("rows in %s is :%d, failed to handle it!!!" % (filename, count))

# workbook.close()

# # Data will be a list of rows represented as dictionaries
# # containing each row's data.
# # data = []
# #
# # keys = None
# # for i, row in enumerate(table.rows):
# #     # text = (cell.text for cell in row.cells)
# #     print("Line %d :" % i)
# #     for cell in row.cells :
# #         print("[%s]" % cell.text)
#     # Establish the mapping based on the first row
#     # headers; these will become the keys of our dictionary
#     # if i == 0:
#     #     keys = tuple(text)
#     #     continue

#     # Construct a dictionary for this row, mapping
#     # keys to values for this row
#     # row_data = dict(zip(keys, text))
#     # data.append(row_data)

# # print(data)
